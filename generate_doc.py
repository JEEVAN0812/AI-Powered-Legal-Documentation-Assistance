from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def add_watermark(document, text):
    # Create a watermark object
    watermark = document.sections[0].footer.paragraphs[0]
    watermark.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    watermark.text = text

def generate_rental_agreement(owner_info, tenant_info, property_info, location, rent_amount):
    doc = Document()

    # Add header/footer
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        run.add_text("AI-Powered Legal Documentation Assistant")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(192, 192, 192)  # Light gray color

    # Title
    doc.add_heading('RENTAL AGREEMENT', level=1)

    # Owner Details
    doc.add_paragraph('Owner:')
    doc.add_paragraph(f'The owner is {owner_info["name"]}, son/daughter of {owner_info["father_name"]}, residing at {owner_info["address"]}.')
    doc.add_paragraph('(hereinafter referred to as the “OWNER”)')

    # Tenant Details
    doc.add_paragraph('\nTenant:')
    doc.add_paragraph(f'The tenant is {tenant_info["name"]}, son/daughter of {tenant_info["father_name"]}, working/studying at {tenant_info["work_address"]}, with a permanent address at {tenant_info["permanent_address"]}.')
    doc.add_paragraph('(hereinafter referred to as the “TENANT”)')

    # Property Details
    doc.add_paragraph('\nProperty Details:')
    doc.add_paragraph(f'The Owner is {owner_info["name"]}, is the absolute owner of the property located at {property_info["address"]}, referred to as the "Demised Premises" in this agreement.')

    # Rent Amount
    doc.add_paragraph('\nRent Amount:')
    doc.add_paragraph(f'The monthly rent amount is Rs. {rent_amount}.')

    # Terms and Conditions
    doc.add_paragraph('\nTerms and Conditions:')
    terms_and_conditions = [
        f"1. The rental period for the Demised Premises shall commence on March 1st, 2024, and continue until February 28th, 2025, with the possibility of extension upon mutual agreement.",
        f"2. The Tenant agrees to pay a monthly rent of Rs. {rent_amount}, excluding electricity and water bills, due on or before the 7th day of each month.",
        # Include other terms and conditions here
    ]
    for term in terms_and_conditions:
        doc.add_paragraph(term)

    # Signature
    doc.add_paragraph('\nIN WITNESS WHEREOF, the parties hereto have executed this agreement as of the date first above written.')
    doc.add_paragraph('[Signature of Owner] _____________________________')
    doc.add_paragraph(owner_info["name"])
    doc.add_paragraph('[Signature of Tenant] _____________________________')
    doc.add_paragraph(tenant_info["name"])

    # Witnesses
    doc.add_paragraph('\nWitnesses:')
    doc.add_paragraph('[Signature of Witness 1] ___________________________')
    doc.add_paragraph('[Signature of Witness 2] ___________________________')
    doc.add_paragraph(f'\nLocation of Witnesses: {location}')

    return doc
